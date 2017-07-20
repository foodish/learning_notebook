
# coding: utf-8

# In[1]:

from docx import Document
from docx.shared import Inches


# In[2]:

document = Document()


# In[3]:

document.add_heading('财务分析_test', 0)


# In[7]:

str_0 = '在已有分析框架的基础上，每次分析时很多部分都是相似的，这就为自动化财务分析提供了机会，这里就进行一下尝试'
p = document.add_paragraph(str_0)


# In[8]:

document.add_heading('公司简介', level=1)


# In[9]:

str1 = '公司是中国白酒龙头，主要生产销售茅台酒及茅台系列酒。历年来茅台酒的销售收入占公司营业收入的90%以上。'
document.add_paragraph(str1)


# In[10]:

document.add_heading('财务分析', level=1)


# In[12]:

document.add_paragraph(
    '基本财务指标', style='ListNumber'
)


# In[13]:

document.add_paragraph('营业收入')


# In[15]:

document.add_picture(r'pic/营收及增长率.png')


# In[16]:

document.add_paragraph(
    '盈利能力分析', style='ListNumber'
)


# In[17]:

document.add_paragraph('净利率')


# In[18]:

document.add_picture(r'pic/净利率对比.png')


# In[19]:

document.save('test.docx')

