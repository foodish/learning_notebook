
# coding: utf-8

# In[1]:

# 引入所需要的库
from docx import Document
from docx.shared import Inches


# In[2]:

# 创建一个docx对象
document = Document()


# In[3]:

# 添加标题
document.add_heading('600519_analysis', 0)


# In[7]:

# 段落中的文字
str_0 = '贵州茅台是白酒龙头，这里进行用其财务数据生成自动化报告'
# 添加段落
p = document.add_paragraph(str_0)


# In[4]:

# 添加1级标题
document.add_heading('公司简介', level=1)


# In[5]:

# 段落中的文字
str1 = '公司是中国白酒龙头，主要生产销售茅台酒及茅台系列酒。历年来茅台酒的销售收入占公司营业收入的90%以上。'
# 添加段落
document.add_paragraph(str1)


# In[6]:

# 添加1级标题
document.add_heading('财务分析', level=1)


# In[7]:

# 添加2级标题
document.add_heading('财务指标分析', level=2)


# In[8]:

# 添加段落
document.add_paragraph('营业收入')


# In[9]:

# 添加此前生成的图片
document.add_picture(r'pic/营业收入.png')


# In[10]:

# 添加段落
document.add_paragraph('归母净利润')


# In[11]:

# 添加此前生成的图片
document.add_picture(r'pic/归母净利润.png')


# In[12]:

# 保存docx并命名
document.save('test.docx')

